"""Tests unitarios puros para app/services/ingestion/parsing/.

Ningún test previo ejercitaba estos parsers directamente (docx.py tenía
~8% de cobertura) - todos los tests de ingestion pasan por el endpoint HTTP
completo, que mockea el parsing. Estos tests llaman las funciones de parsing
directamente con archivos reales (generados con python-docx / pandas) y no
requieren client/db_session ni MySQL.
"""
from __future__ import annotations

import pytest

from app.services.ingestion.parsing.docx import parse_docx
from app.services.ingestion.parsing.txt import parse_txt
from app.services.ingestion.parsing.pdf import parse_pdf
from app.services.ingestion.parsing.dispatcher import parse_source
from app.models.enums import SourceType


# ── docx.py ──────────────────────────────────────────────────────────────

def _add_numbered_paragraph(doc, text: str, ilvl: int = 0) -> None:
    """python-docx no agrega w:numPr real solo con style='List Bullet'
    (el estilo por sí solo no basta - Word decide "es lista" por la
    presencia de w:numPr en w:pPr). Lo inyectamos a mano para simular
    una lista auto-numerada real de Word, que es lo que _detect_sections
    / parse_docx buscan."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    para = doc.add_paragraph(text)
    pPr = para._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), "1")
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)


def _make_docx(path: str) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("Título Principal", level=0)  # style "Title"
    doc.add_heading("Sección Uno", level=1)  # "Heading 1"
    doc.add_paragraph("Este es un párrafo normal.")

    bold_para = doc.add_paragraph()
    bold_run = bold_para.add_run("Texto en negrita completo")
    bold_run.bold = True

    _add_numbered_paragraph(doc, "Ítem de lista 1")
    _add_numbered_paragraph(doc, "Ítem de lista 2", ilvl=1)

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"

    doc.add_paragraph("")  # párrafo vacío: debe ser omitido

    doc.save(path)


class TestParseDocx:
    async def test_extracts_heading_as_markdown(self, tmp_path):
        path = str(tmp_path / "doc.docx")
        _make_docx(path)

        text = await parse_docx(path)

        assert "# Título Principal" in text
        assert "## Sección Uno" in text

    async def test_extracts_plain_paragraph(self, tmp_path):
        path = str(tmp_path / "doc.docx")
        _make_docx(path)

        text = await parse_docx(path)

        assert "Este es un párrafo normal." in text

    async def test_bold_paragraph_becomes_markdown_bold(self, tmp_path):
        path = str(tmp_path / "doc.docx")
        _make_docx(path)

        text = await parse_docx(path)

        assert "**Texto en negrita completo**" in text

    async def test_list_items_become_markdown_list(self, tmp_path):
        path = str(tmp_path / "doc.docx")
        _make_docx(path)

        text = await parse_docx(path)

        assert "- Ítem de lista 1" in text
        assert "  - Ítem de lista 2" in text  # ilvl=1 → sangría de 2 espacios

    async def test_bold_list_item_becomes_bold_markdown_list(self, tmp_path):
        from docx import Document

        path = str(tmp_path / "bold_list.docx")
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Ítem en negrita")
        run.bold = True
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        pPr = para._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        ilvl_el = OxmlElement("w:ilvl")
        ilvl_el.set(qn("w:val"), "0")
        numId_el = OxmlElement("w:numId")
        numId_el.set(qn("w:val"), "1")
        numPr.append(ilvl_el)
        numPr.append(numId_el)
        pPr.append(numPr)
        doc.save(path)

        text = await parse_docx(path)

        assert "**Ítem en negrita**" in text
        assert "- **Ítem en negrita**" not in text

    async def test_table_becomes_markdown_table(self, tmp_path):
        path = str(tmp_path / "doc.docx")
        _make_docx(path)

        text = await parse_docx(path)

        assert "| A | B |" in text
        assert "|---|---|" in text
        assert "| 1 | 2 |" in text

    async def test_empty_paragraphs_are_skipped(self, tmp_path):
        path = str(tmp_path / "doc.docx")
        _make_docx(path)

        text = await parse_docx(path)

        # No debe haber tres o más saltos de línea seguidos por un párrafo vacío colado
        assert "\n\n\n" not in text

    async def test_heading_es_variant_is_recognized(self, tmp_path):
        from docx import Document

        path = str(tmp_path / "doc_es.docx")
        doc = Document()
        para = doc.add_paragraph("Encabezado en español")
        para.style = doc.styles["Heading 1"]
        # Renombrar el estilo asignado para simular variante localizada
        # (python-docx no permite crear estilos "Título 1" fácilmente sin
        # una plantilla localizada, así que verificamos vía el diccionario
        # _HEADING_STYLES directamente en un test aparte)
        doc.save(path)

        text = await parse_docx(path)
        assert "## Encabezado en español" in text

    def test_heading_styles_dict_covers_es_variants(self):
        from app.services.ingestion.parsing.docx import _HEADING_STYLES

        assert _HEADING_STYLES["título"] == "#"
        assert _HEADING_STYLES["encabezado 1"] == "##"
        assert _HEADING_STYLES["título 2"] == "###"
        assert _HEADING_STYLES["encabezado 4"] == "#####"

    async def test_nonexistent_file_raises_runtime_error(self, tmp_path):
        path = str(tmp_path / "does_not_exist.docx")

        with pytest.raises(RuntimeError, match="No se pudo parsear el DOCX"):
            await parse_docx(path)

    async def test_corrupt_file_raises_runtime_error(self, tmp_path):
        path = str(tmp_path / "corrupt.docx")
        with open(path, "wb") as f:
            f.write(b"not a real docx file")

        with pytest.raises(RuntimeError, match="No se pudo parsear el DOCX"):
            await parse_docx(path)


class TestZipBombProtection:
    """python-docx descomprime el ZIP interno sin validar el ratio de
    expansión - un DOCX construido con XML muy repetitivo puede expandirse
    órdenes de magnitud más que su tamaño en disco. _check_zip_bomb valida
    esto antes de invocar Document()."""

    async def test_normal_docx_passes(self, tmp_path):
        path = str(tmp_path / "doc.docx")
        _make_docx(path)
        # No debe lanzar - un DOCX normal generado por python-docx tiene
        # una tasa de compresión típica de texto/XML, muy por debajo del umbral.
        text = await parse_docx(path)
        assert text  # confirma que sí llegó a parsear, no solo que no lanzó

    async def test_high_compression_ratio_is_rejected(self, tmp_path):
        import zipfile
        from app.services.ingestion.parsing.docx import _check_zip_bomb

        path = str(tmp_path / "bomb.docx")
        # Un solo miembro con contenido extremadamente repetitivo comprime a
        # una fracción mínima de su tamaño real - simula el patrón de una
        # zip bomb sin necesitar gigabytes reales en disco de test.
        payload = b"A" * (50 * 1024 * 1024)  # 50MB de un solo byte repetido
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED, compresslevel=9) as zf:
            zf.writestr("word/document.xml", payload)

        with pytest.raises(ValueError, match="tasa de compresión"):
            _check_zip_bomb(path)

    async def test_within_normal_ratio_passes(self, tmp_path):
        import os
        import zipfile
        from app.services.ingestion.parsing.docx import _check_zip_bomb

        path = str(tmp_path / "normal.docx")
        # os.urandom no comprime bien (alta entropía, como el contenido real
        # de un DOCX con texto variado + metadata XML) - a diferencia de un
        # payload repetitivo, que comprime órdenes de magnitud mejor y por
        # eso no sirve como "caso normal" para este test.
        payload = os.urandom(200 * 1024)
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("word/document.xml", payload)

        _check_zip_bomb(path)  # no debe lanzar


# ── txt.py ───────────────────────────────────────────────────────────────

class TestParseTxt:
    async def test_reads_utf8_file(self, tmp_path):
        path = tmp_path / "file.txt"
        path.write_text("Hola  mundo con acentos: ñ, á, é", encoding="utf-8")

        text = await parse_txt(str(path))

        assert text == "Hola  mundo con acentos: ñ, á, é"

    async def test_strips_leading_trailing_whitespace(self, tmp_path):
        path = tmp_path / "file.txt"
        path.write_text("\n\n  contenido  \n\n", encoding="utf-8")

        text = await parse_txt(str(path))

        assert text == "contenido"

    async def test_falls_back_to_latin1_when_not_utf8(self, tmp_path):
        path = tmp_path / "file.txt"
        # 'ñ' en latin-1 no es UTF-8 válido en ese byte
        path.write_bytes("año de la señorita".encode("latin-1"))

        text = await parse_txt(str(path))

        assert "año de la señorita" == text

    async def test_missing_file_raises(self, tmp_path):
        path = tmp_path / "missing.txt"

        with pytest.raises(FileNotFoundError):
            await parse_txt(str(path))

    async def test_binary_content_raises_instead_of_silently_parsing_garbage(self, tmp_path):
        """latin-1 decodifica CUALQUIER byte sin UnicodeDecodeError (mapea
        1:1 byte->codepoint), así que antes de esta validación un archivo
        binario renombrado a .txt (p. ej. un .docx real) 'parseaba' sin
        error como mojibake ilegible - producía chunks y embeddings de
        basura que quedaban status=ready sin ninguna alerta de calidad."""
        path = tmp_path / "file.txt"
        # Bytes de control aleatorios, no texto real en ningún encoding.
        path.write_bytes(bytes(range(0, 40)) * 50)

        with pytest.raises(RuntimeError, match="no parece texto legible"):
            await parse_txt(str(path))


# ── pdf.py ───────────────────────────────────────────────────────────────

class TestParsePdf:
    """pdf.py fue reescrito (commit 4014b6f) para usar pymupdf4llm.to_markdown
    en dos intentos (force_text=True, luego force_text=False para OCR) en
    vez de las funciones _extract_pymupdf/_extract_ocr que estos tests
    mockeaban antes - quedaron probando una API que ya no existe
    (ImportError en collection). Reescritos contra la implementación real.

    pdf.py hace `import pymupdf4llm` DENTRO de parse_pdf() (no a nivel de
    módulo), así que el mock debe aplicarse sobre el paquete real
    (pymupdf4llm.to_markdown) vía monkeypatch - parchear un atributo en
    pdf_mod no tendría efecto porque el nombre no existe ahí hasta que la
    función corre, y el import local siempre trae el real.
    """

    async def test_uses_force_text_result_when_present(self, monkeypatch):
        import pymupdf4llm
        from app.services.ingestion.parsing import pdf as pdf_mod

        calls = []

        def fake_to_markdown(path, force_text=True):
            calls.append(force_text)
            return "x" * 500

        monkeypatch.setattr(pymupdf4llm, "to_markdown", fake_to_markdown)

        text = await pdf_mod.parse_pdf("fake.pdf")

        assert text == "x" * 500
        assert calls == [True]  # no debió intentar el fallback OCR

    async def test_falls_back_to_ocr_when_force_text_returns_empty(self, monkeypatch):
        import pymupdf4llm
        from app.services.ingestion.parsing import pdf as pdf_mod

        def fake_to_markdown(path, force_text=True):
            return "" if force_text else "contenido ocr"

        monkeypatch.setattr(pymupdf4llm, "to_markdown", fake_to_markdown)

        text = await pdf_mod.parse_pdf("fake.pdf")

        assert text == "contenido ocr"

    async def test_falls_back_to_ocr_when_force_text_raises(self, monkeypatch):
        import pymupdf4llm
        from app.services.ingestion.parsing import pdf as pdf_mod

        def fake_to_markdown(path, force_text=True):
            if force_text:
                raise ValueError("corrupt pdf")
            return "contenido ocr"

        monkeypatch.setattr(pymupdf4llm, "to_markdown", fake_to_markdown)

        text = await pdf_mod.parse_pdf("fake.pdf")

        assert text == "contenido ocr"

    async def test_raises_runtime_error_when_both_strategies_fail(self, monkeypatch):
        import pymupdf4llm
        from app.services.ingestion.parsing import pdf as pdf_mod

        def boom(path, force_text=True):
            raise ValueError("corrupt pdf")

        monkeypatch.setattr(pymupdf4llm, "to_markdown", boom)

        with pytest.raises(RuntimeError, match="No se pudo parsear el PDF"):
            await pdf_mod.parse_pdf("fake.pdf")

    async def test_parses_real_pdf_with_text(self, tmp_path):
        """Ejercita parse_pdf real (sin mock) contra un PDF generado con pymupdf."""
        import pymupdf

        path = str(tmp_path / "full.pdf")
        doc = pymupdf.open()
        page = doc.new_page()
        long_text = "Lorem ipsum dolor sit amet. " * 20
        page.insert_text((72, 72), long_text)
        doc.save(path)
        doc.close()

        from app.services.ingestion.parsing.pdf import parse_pdf
        result = await parse_pdf(path)

        assert "Lorem ipsum" in result


# ── dispatcher.py ────────────────────────────────────────────────────────

class TestDispatcher:
    async def test_dispatches_txt(self, tmp_path):
        path = tmp_path / "file.txt"
        path.write_text("contenido de prueba", encoding="utf-8")

        text = await parse_source(SourceType.txt, str(path))

        assert text == "contenido de prueba"

    async def test_unsupported_source_type_raises_value_error(self):
        with pytest.raises(ValueError, match="Tipo de fuente no soportado"):
            await parse_source(SourceType.faq, "whatever.faq")

    async def test_missing_file_path_raises_value_error(self):
        with pytest.raises(ValueError, match="requiere file_path"):
            await parse_source(SourceType.txt, None)

        with pytest.raises(ValueError, match="requiere file_path"):
            await parse_source(SourceType.txt, "")
